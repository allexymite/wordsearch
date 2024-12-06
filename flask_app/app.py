# TODO
# # finish security updates from chatgpt

from flask import Flask, render_template, request, send_file
import os
import random
import string
from PIL import Image, ImageDraw, ImageFont
from docx import Document

app = Flask(__name__)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = 'uploads'
STATIC_FOLDER = os.path.join(BASE_DIR, "static")
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# validate input
def validate_input(words):
    """Validate the words input to ensure it's safe and meets expected format."""
    if not words:
        return False
    # Allow only alphabetic characters and commas (modify as needed for your use case)
    for word in words.split(","):
        word = word.strip()
        if not word.isalpha():
            return False
    return True


# making the word search
def generate_word_search(words, size=15):
    # Create an empty grid
    grid = [[" " for _ in range(size)] for _ in range(size)]
    
    def place_word(word):
        word = word.upper()
        word_len = len(word)
        orientations = [(1, 0), (0, 1), (1, 1), (1, -1)]  # Horizontal, Vertical, Diagonal (2 directions)
        random.shuffle(orientations)

        for _ in range(100):  # Try 100 random placements
            direction = random.choice(orientations)
            start_row = random.randint(0, size - 1)
            start_col = random.randint(0, size - 1)
            end_row = start_row + direction[0] * (word_len - 1)
            end_col = start_col + direction[1] * (word_len - 1)

            # Check boundaries
            if not (0 <= end_row < size and 0 <= end_col < size):
                continue

            # Check for overlap
            can_place = True
            row, col = start_row, start_col
            for letter in word:
                if grid[row][col] not in (" ", letter):
                    can_place = False
                    break
                row += direction[0]
                col += direction[1]

            if can_place:
                # Place the word
                row, col = start_row, start_col
                for letter in word:
                    grid[row][col] = letter
                    row += direction[0]
                    col += direction[1]
                return True
        return False

    # Place all words in the grid
    for word in words:
        if not place_word(word):
            print(f"Could not place the word: {word}")

    # Fill remaining spaces with random letters
    for row in range(size):
        for col in range(size):
            if grid[row][col] == " ":
                grid[row][col] = random.choice(string.ascii_uppercase)

    return grid

# convert to docx
def save_grid_as_docx(grid, words, filename="word_search.docx"):
    doc = Document()
    doc.add_heading("Word Search Puzzle", level=1)

    # Add grid
    table = doc.add_table(rows=len(grid), cols=len(grid[0]))
    for i, row in enumerate(grid):
        for j, letter in enumerate(row):
            cell = table.cell(i, j)
            cell.text = letter
            cell.width = 300000  # Adjust cell width as needed

    # Add word list
    doc.add_paragraph("\nWords to Find:")
    for word in words:
        doc.add_paragraph(word.upper())

    doc.save(filename)
    print(f"Saved to {filename}")

# save grid as image
def save_grid_as_image(grid, output_path):
    cell_size = 50
    image_size = len(grid) * cell_size
    image = Image.new("RGB", (image_size, image_size), "white")
    draw = ImageDraw.Draw(image)
    
    # Load a truetype font
    font_path = os.path.join(BASE_DIR, "fonts", "Lexend-Medium.ttf")

    font_size = cell_size // 2  # Adjust font size relative to cell size
    font = ImageFont.truetype(font_path, font_size)


    for i, row in enumerate(grid):
        for j, letter in enumerate(row):
            x = j * cell_size + cell_size // 4
            y = i * cell_size + cell_size // 4
            draw.text((x, y), letter, fill="black", font=font)
            # draw.rectangle(
            #     [j * cell_size, i * cell_size, (j + 1) * cell_size, (i + 1) * cell_size],
            #     outline="black",
            # )

    image.save(output_path)


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        words = request.form.get("words")
        if "file" in request.files and request.files["file"].filename != "":
            file = request.files["file"]
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(filepath)
            with open(filepath, "r") as f:
                words = f.read()
        if not words:
            return render_template("index.html", error="Please provide words.")
        
        words = [word.strip() for word in words.split(",")]

        # Generate Word Search
        grid = generate_word_search(words)
        image_path = os.path.join(STATIC_FOLDER, "word_search_image.png")
        save_grid_as_image(grid, image_path)

        # Save DOCX
        docx_path = os.path.join(STATIC_FOLDER, "word_search.docx")
        save_grid_as_docx(grid, words, docx_path)

        return render_template(
            "index.html", image_path=image_path, docx_path=docx_path, success=True
        )
    return render_template("index.html")


@app.route("/download/<path:filename>")
def download_file(filename):
    return send_file(os.path.join(STATIC_FOLDER, filename), as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True)
